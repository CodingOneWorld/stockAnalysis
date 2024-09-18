# -*- coding: utf-8 -*-
# 判断股票的关键价位
# 如支撑位，压力位，均线位置等
# 计算当前价格最高值，最低值，以及均线数据，进行比较
from analysis_util.cal_key_price import cal_extreme_min_value
from analysis_util.cal_mean_line import cal_all_mean_line
from analysis_util.cal_stock_trend import get_stock_price, cal_stock_price_trend, cal_trend_common
from analysis_util.plot_k_line import plot_k_line, plot_k_line_latestdays, save_k_line

from docx.shared import Cm
from analysis_util.output_document import Doc

import pandas as pd
import numpy as np


def output_doc(df, file_path):
    # doc文档
    doc = Doc()
    for line in df.values:
        print(line)
        doc.add_heading('，'.join(line))
        symbol = line[0]

        # 画出其最近100天，300天，1000天日线图
        save_k_line(symbol, 100, './resources/kline100.png')
        save_k_line(symbol, 300, './resources/kline300.png')
        save_k_line(symbol, 1000, './resources/kline1000.png')
        doc.add_picture('./resources/kline100.png', width=Cm(10))
        doc.add_picture('./resources/kline300.png', width=Cm(10))
        doc.add_picture('./resources/kline1000.png', width=Cm(10))

    # 保存文档
    doc.save(file_path)


# 指定均线进行比较
# 均线应呈上升趋势，然后股价呈下降趋势，此时存在支撑
def mav_compare(code, mav_list):
    '''
    :param code: 股票代码 6位数
    :param mav_list:  需要比较的均线列表
    :return:
    '''
    # 获取股票交易数据
    his_price_df = get_stock_price(code, 'close')[-300:]
    his_price = his_price_df['close'].values
    print(his_price)
    # 计算均线
    stock_mean_list = []
    for mav in mav_list:
        his_price_df['mean' + str(mav)] = his_price_df.close.rolling(window=mav).mean().fillna(0)
        print(his_price_df.head())
        # 均线应呈上升趋势
        k=cal_trend_common(his_price_df['mean' + str(mav)][-mav:].values)
        # 当前价格与均线进行比较
        cur_price = his_price[-1]
        mean = his_price_df['mean' + str(mav)].values[-1]
        print(cur_price, mean)
        if mean * 0.95 <= cur_price <= mean * 1.05 and k>0:
            stock_mean_list.append(mav)
    print(stock_mean_list)
    if len(stock_mean_list) > 0:
        print('%s接近以下均线%s' % (code, ','.join([str(i) for i in stock_mean_list])))

    return stock_mean_list


# 计算全部均线进行比较
def all_mean_price_compare(doc, stock):
    # doc = Doc()

    code = stock[0]
    # 计算股票趋势
    stock_price = get_stock_price(code, 'close')
    k = 1 if cal_stock_price_trend(stock_price, 10) > 0 else -1
    print(k)

    if k < 0:
        # 获取股票最近的价格
        stock_price = get_stock_price(code, 'low').values[-1]

        # 计算均线值
        mean_dict = cal_all_mean_line(code, 3000)
        print(mean_dict)

        # 股价与均线值比较
        stock_mean_list = []
        for mean in mean_dict.keys():
            if mean_dict.get(mean) * 0.95 <= stock_price <= mean_dict.get(mean) * 1.05:
                stock_mean_list.append(mean)

        doc.add_heading('，'.join(','.join(stock)))
        if len(stock_mean_list) > 0:
            if len(stock_mean_list) > 0:
                print(code + ": " + "接近以下均线: " + ','.join(stock_mean_list))
                doc.add_heading("接近以下均线: " + ','.join(stock_mean_list))

            # 画出其最近100天，300天，1000天日线图
            save_k_line(code, 100, './resources/kline100.png')
            save_k_line(code, 300, './resources/kline300.png')
            doc.add_picture('./resources/kline100.png', width=Cm(10))
            doc.add_picture('./resources/kline300.png', width=Cm(10))

    return 0


def extreme_price_compare(code):
    # 与历史股价极小值比较
    # 获取股票历史价格
    his_price = get_stock_price(code, 'low')['low'].values[-300:]
    # print(his_price)
    current_price = his_price[-1]
    # 计算历史股价极小值
    ex_price = cal_extreme_min_value(his_price)
    ex_price = cal_extreme_min_value(ex_price)

    ex_price_list = []
    for price in ex_price:
        if price * 0.95 <= current_price <= price * 1.05:
            ex_price_list.append(price)

    if len(ex_price_list) > 0:
        print(code + ": " + "接近以下历史股价低点: " + ','.join([str(i) for i in ex_price_list]))


if __name__ == '__main__':
    # 获取自选股票池
    # file = '自选股.csv'
    file = 'stock_pool2023.txt'
    df = pd.read_csv(file, dtype={'symbol': np.str}, delimiter=',')
    stock_list = df.values
    selected_stock = []
    for s in stock_list:
        print(s)
        his_price = get_stock_price(s[0], 'low')['low'].values[-300:]
        k10 = cal_stock_price_trend(his_price, 10)
        k100 = cal_stock_price_trend(his_price, 100)
        stock_mean_list = mav_compare(s[0], [60, 140])
        if len(stock_mean_list) > 0 and k10 < -0.1 and k100 < 0:
            res = '%s接近以下均线%s' % (s[0], ','.join([str(i) for i in stock_mean_list]))
            selected_stock.append([s[0], s[1], res])

    # 保存文档
    path = '关键价格及均线比较.docx'
    df = pd.DataFrame(selected_stock, columns=['code', 'name', 'mean'])
    output_doc(df, path)

    # mav_compare('002028', [60])

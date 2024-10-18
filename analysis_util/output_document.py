# -*- coding: utf-8 -*-

from docx import Document
# docx.shared 用于设置大小（图片等）
from docx.shared import Cm, Pt
from docx.document import Document as Doc

# 将分析结果输出到word文档中
from analysis_util.plot_k_line import save_k_line


def Doc():
    # 创建代表Word文档的Doc对象
    document = Document()
    return document


def output_doc(df, file_path):
    '''

    :param df: DataFrame columns=['symbol','stock_name']
    :param file_path: 输出文档的路径
    :return:
    '''
    # doc文档
    doc = Doc()
    for line in df.values:
        print(line)
        doc.add_heading('，'.join(line))
        symbol = line[0]

        # 添加段落
        # if len(paragraph)>0:
        #     doc.add_paragraph(paragraph)

        # 画出其最近100天，300天，1000天日线图
        save_k_line(symbol, 100, './resources/kline100.png')
        save_k_line(symbol, 300, './resources/kline300.png')
        save_k_line(symbol, 1000, './resources/kline1000.png')
        doc.add_picture('./resources/kline100.png', width=Cm(10))
        doc.add_picture('./resources/kline300.png', width=Cm(10))
        doc.add_picture('./resources/kline1000.png', width=Cm(10))

    # 保存文档
    doc.save(file_path)


def example():
    # 创建代表Word文档的Doc对象
    document = Document()
    # type: Doc
    # 添加大标
    document.add_heading('自', 0)
    # 添加段落
    p = document.add_paragraph('Python是一门非常流行的编程语言')
    run = p.add_run('very easy')
    run.bold = True
    run.font.size = Pt(18)
    p.add_run('hello')
    run = p.add_run('非常棒')
    run.font.size = Pt(18)
    run.underline = False
    p.add_run('。')
    # 添加一级标题
    document.add_heading('Heading, level 1', level=1)
    # 添加带样式的段落
    document.add_paragraph('Intense quote', style='Intense Quote')
    # 添加无序列表
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'second item in ordered list', style='List Bullet'
    )
    # 添加有序列表
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    document.add_paragraph(
        'second item in ordered list', style='List Number'
    )
    # 添加图片（注意路径和图片必须要存在）
    # document.add_picture('a.jpg', width=Cm(5.2))

    # 添加分节符
    document.add_section()
    records = (
        ('亚瑟', '战士英雄'),
        ('白起', '坦克英雄'),
        ('赵云', '刺客英雄'),
        ('女娲', '法师英雄'),
    )

    # 添加表格
    table = document.add_table(rows=1, cols=3)
    table.style = 'Dark List'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ' 姓 名 '
    hdr_cells[1].text = '类别'

    # 为表格添加行
    for name, sex in records:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = sex
        # row_cells[2].text = birthday
    # 添加分页符
    document.add_page_break()
    # 保存文档
    document.save('demo.docx')


if __name__ == '__main__':
    example()

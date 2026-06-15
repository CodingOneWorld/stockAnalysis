# -*- coding: utf-8 -*-
"""
将"基本面好的股票.txt"中的每只股票近一年 K 线图批量导出到 Word 文档。

功能：
  - 读取 stock_select/基本面好的股票.txt
  - 逐只从本地数据库获取近 252 个交易日（约一年）的日线数据
  - 用 mplfinance 生成K线图（含均线 5/10/20/60 + 成交量），A 股配色（红涨绿跌）
  - 将所有 K 线图按每页两列插入 Word 文档
  - 输出：stock_select/基本面好的股票_K线图.docx

运行方式（两种均可）：
  cd /Users/zhangqi21/PyCharmMiscProject/stockAnalysis
  python -m stock_select.export_kline_to_doc
  # 或直接在 PyCharm 中右键运行本文件
"""
import io
import datetime
import pathlib
import sys
import os

# ── 确保项目根目录在 sys.path，兼容直接运行和 -m 两种方式 ──────────────
_THIS_DIR     = pathlib.Path(__file__).resolve().parent   # stock_select/
_PROJECT_ROOT = _THIS_DIR.parent                          # 项目根目录
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

import pandas as pd
import mplfinance as mpf
import matplotlib
matplotlib.use('Agg')   # 非交互后端，避免弹窗
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from trade_data.get_trade_data import get_stock_trade_data

# ── 配置（路径均基于脚本所在目录，与工作目录无关）──────────────────────
STOCK_LIST_FILE = str(_THIS_DIR / '基本面好的股票.txt')
OUTPUT_DOC      = str(_THIS_DIR / '基本面好的股票_K线图.docx')
LATEST_DAYS     = 252          # 近一年（约252个交易日）
MAV             = (5, 10, 20, 60)
COLS_PER_ROW    = 2            # Word 表格每行放几张图
IMG_WIDTH_INCH  = 3.2          # 每张图在 Word 中的宽度（英寸）
FIG_SIZE        = (10, 6)      # mplfinance 输出图像尺寸（英寸）


# ── A股配色：红涨绿跌 ─────────────────────────────────────────────────
_MC = mpf.make_marketcolors(
    up='red', down='green',
    edge='inherit', wick='inherit', volume='inherit'
)
_STYLE = mpf.make_mpf_style(
    marketcolors=_MC,
    gridcolor='#e0e0e0',
    rc={'font.family': 'Heiti TC',   # macOS 黑体-简，支持中文
        'axes.unicode_minus': False}
)


def _load_kline(symbol: str, latest_days: int = LATEST_DAYS) -> pd.DataFrame | None:
    """从数据库读取近 latest_days 个交易日数据，整理为 mplfinance 所需格式。"""
    df = get_stock_trade_data(symbol)
    if df is None or df.empty:
        return None

    df = df.tail(latest_days).copy()
    df['trade_date'] = pd.to_datetime(df['trade_date'].astype(str), format='%Y%m%d')
    df = df.set_index('trade_date')
    df = df[['open', 'high', 'close', 'low', 'vol']].rename(columns={'vol': 'volume'})
    df = df.dropna()
    return df


def _render_kline_to_bytes(symbol: str, name: str, df: pd.DataFrame) -> bytes:
    """用 mplfinance 渲染K线图，返回 PNG 字节流。"""
    title = f'{name}（{symbol}）  近一年K线'
    fig, axes = mpf.plot(
        df,
        type='candle',
        style=_STYLE,
        mav=MAV,
        volume=True,
        title=title,
        figsize=FIG_SIZE,
        returnfig=True,
    )
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def export_to_word(stock_list_file: str = STOCK_LIST_FILE,
                   output_doc: str = OUTPUT_DOC,
                   latest_days: int = LATEST_DAYS,
                   cols: int = COLS_PER_ROW):
    """主函数：批量生成K线图并写入 Word 文档。"""
    # ── 读取股票列表 ──────────────────────────────────────────────────
    stocks = pd.read_csv(stock_list_file, dtype={'symbol': str})
    total = len(stocks)
    print(f'[导出] 共 {total} 只股票，开始生成K线图...')

    # ── 初始化 Word 文档 ───────────────────────────────────────────────
    doc = Document()

    # 页面设置：A4 横向，边距收窄
    section = doc.sections[0]
    section.page_width  = int(11906 * 914)   # A4宽
    section.page_height = int(16838 * 914)   # A4高（竖向）
    # 上下左右边距 1.5cm
    from docx.shared import Cm
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # 标题
    title_para = doc.add_heading('基本面优质股票 · 近一年K线图', level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 副标题：生成时间
    sub = doc.add_paragraph(
        f'生成时间：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}  '
        f'  均线：{"、".join(str(m)+"日" for m in MAV)}  '
        f'  数据范围：近 {latest_days} 个交易日'
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()  # 空行

    # ── 逐只渲染并填入表格 ────────────────────────────────────────────
    # 按 cols 列布局，每 cols 只股票一行，用 Word 表格对齐
    rows_needed = (total + cols - 1) // cols
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.style = 'Table Grid'

    # 隐藏表格边框
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    def _hide_borders(tbl):
        tbl_pr = tbl._tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl._tbl.insert(0, tbl_pr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            borders.append(el)
        tbl_pr.append(borders)
    _hide_borders(table)

    success_count = 0
    skip_count    = 0

    for idx, row_data in stocks.iterrows():
        symbol = str(row_data['symbol']).zfill(6)
        name   = str(row_data['stock_name'])

        print(f'  [{idx + 1}/{total}] {symbol} {name} ...', end=' ')

        df = _load_kline(symbol, latest_days)
        if df is None or len(df) < 20:
            print('⚠ 数据不足，跳过')
            skip_count += 1
            # 填入空格占位，保持表格对齐
            r_idx = idx // cols
            c_idx = idx % cols
            table.cell(r_idx, c_idx).text = f'{symbol} {name}\n（数据不足）'
            continue

        try:
            img_bytes = _render_kline_to_bytes(symbol, name, df)
        except Exception as e:
            print(f'✗ 渲染失败: {e}')
            skip_count += 1
            r_idx = idx // cols
            c_idx = idx % cols
            table.cell(r_idx, c_idx).text = f'{symbol} {name}\n（渲染失败）'
            continue

        r_idx = idx // cols
        c_idx = idx % cols
        cell  = table.cell(r_idx, c_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 股票标题
        p_title = cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(f'{name}（{symbol}）')
        run_title.bold = True
        run_title.font.size = Pt(9)

        # 插入图片
        p_img = cell.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(IMG_WIDTH_INCH))

        print('✓')
        success_count += 1

    # 补全最后一行空白单元格（cols不整除时）
    remainder = total % cols
    if remainder != 0:
        last_row = rows_needed - 1
        for c_idx in range(remainder, cols):
            table.cell(last_row, c_idx).text = ''

    # ── 保存 ──────────────────────────────────────────────────────────
    doc.save(output_doc)
    print(f'\n[导出] 完成！成功 {success_count} 只，跳过 {skip_count} 只')
    print(f'[导出] 文档已保存：{output_doc}')


if __name__ == '__main__':
    export_to_word()
